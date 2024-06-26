import warnings
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from data.fetchers import fetch_match_data, fetch_player_data
from players.player_tables import create_player_tables

# Suppress specific deprecation warnings
warnings.filterwarnings('ignore', category=UserWarning,
                        message='.*style lookup by style_id is deprecated.*')


def create_full_report():
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('NRL Full Report', level=0)

    doc.add_paragraph()

    # Add a centered image
    para_image = doc.add_paragraph()
    run_image = para_image.add_run()
    run_image.add_picture('images/stock_image.jpg', width=Cm(12))
    para_image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the paragraph

    # add a page
    doc.add_page_break()

    doc.add_paragraph('The History of NRL', style='Heading 1')

    # add content to the second page
    history_text = (
        "A full report on the team (limited data) and selected player based on ChatGPT"
    )

    doc.add_paragraph(history_text, style='BodyText')

    # Align paragraph
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # add a page
    doc.add_page_break()

    doc.add_paragraph('NRL MATCH DATA:', style='Heading 1')

    fetch_match_data()

    # add a page
    doc.add_page_break()

    doc.add_paragraph('NRL PLAYER DATA:', style='Heading 1')
    df_player = fetch_player_data()
    if df_player is not None:
        create_player_tables(df_player, doc)

    # allocate save name, save location, save doc and open doc
    output_folder = 'report_outputs'
    filename = f"{output_folder}/full_report_test.docx"

    # Save the document
    doc.save(filename)

    # Open the doc with the default application
    os.system(f"open '{filename}'")
